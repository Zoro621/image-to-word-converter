"""
Word document (.docx) generation from extracted text with formatting.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from io import BytesIO


class DocxGenerator:
    """
    Generate formatted Word documents from extracted text.
    """
    
    def __init__(self):
        """Initialize the document generator."""
        self.document = None
    
    def create_document(self, extracted_text: str, title: str = "Extracted Document") -> BytesIO:
        """
        Create a Word document from extracted text.
        
        Args:
            extracted_text: Text with formatting markers from vision model
            title: Document title
            
        Returns:
            BytesIO object containing the .docx file
        """
        self.document = Document()
        
        # Set up document styles
        self._setup_styles()
        
        # Add title
        title_para = self.document.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a subtle separator
        self._add_separator()
        
        # Process and add content
        self._process_content(extracted_text)
        
        # Save to BytesIO
        doc_buffer = BytesIO()
        self.document.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
    
    def _setup_styles(self):
        """Set up custom document styles."""
        styles = self.document.styles
        
        # Customize Normal style
        style = styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Customize Heading styles
        for i in range(1, 4):
            style = styles[f'Heading {i}']
            font = style.font
            font.name = 'Calibri'
            font.bold = True
            font.color.rgb = RGBColor(30, 64, 175)  # Deep blue
    
    def _add_separator(self):
        """Add a horizontal separator line."""
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('─' * 50)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(8)
    
    def _process_content(self, text: str):
        """Process extracted text and add formatted content to document."""
        lines = text.split('\n')
        current_list_type = None
        list_items = []
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Empty line - process any pending list
                if list_items:
                    self._add_list(list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                i += 1
                continue
            
            # Main heading (##)
            if line.startswith('## ') and not line.startswith('### '):
                if list_items:
                    self._add_list(list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                self._add_heading(line[3:], level=1)
            
            # Subheading (###)
            elif line.startswith('### '):
                if list_items:
                    self._add_list(list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                self._add_heading(line[4:], level=2)
            
            # Top-level heading (#)
            elif line.startswith('# ') and not line.startswith('## '):
                if list_items:
                    self._add_list(list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                self._add_heading(line[2:], level=0)
            
            # Bullet point
            elif line.startswith('- ') or line.startswith('• '):
                current_list_type = 'bullet'
                list_items.append(line[2:])
            
            # Numbered list
            elif re.match(r'^\d+\.\s', line):
                current_list_type = 'numbered'
                list_items.append(re.sub(r'^\d+\.\s', '', line))
            
            # Blockquote
            elif line.startswith('> '):
                if list_items:
                    self._add_list(list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                self._add_blockquote(line[2:])
            
            # Diagram marker
            elif '[DIAGRAM:' in line:
                if list_items:
                    self._add_list(list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                diagram_match = re.search(r'\[DIAGRAM:\s*(.*?)\]', line)
                if diagram_match:
                    self._add_diagram_placeholder(diagram_match.group(1))
            
            # Table (detect by pipe characters)
            elif '|' in line and line.count('|') >= 2:
                if list_items:
                    self._add_list(list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                # Collect table lines
                table_lines = [line]
                j = i + 1
                while j < len(lines) and '|' in lines[j]:
                    table_lines.append(lines[j].strip())
                    j += 1
                self._add_table_from_markdown(table_lines)
                i = j - 1
            
            # Regular paragraph
            else:
                if list_items:
                    self._add_list(list_items, current_list_type)
                    list_items = []
                    current_list_type = None
                self._add_paragraph(line)
            
            i += 1
        
        # Process any remaining list items
        if list_items:
            self._add_list(list_items, current_list_type)
    
    def _add_heading(self, text: str, level: int):
        """Add a heading to the document."""
        heading = self.document.add_heading(text, level=level)
    
    def _add_paragraph(self, text: str):
        """Add a paragraph with inline formatting."""
        paragraph = self.document.add_paragraph()
        self._add_formatted_text(paragraph, text)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text to a paragraph with bold/italic formatting."""
        # Pattern to match **bold** and *italic*
        pattern = r'(\*\*.*?\*\*|\*[^*]+\*|\$.*?\$)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('$') and part.endswith('$'):
                # Formula/equation - use italic and different color
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                run.font.color.rgb = RGBColor(139, 92, 246)  # Purple for formulas
            else:
                # Regular text
                paragraph.add_run(part)
    
    def _add_list(self, items: list, list_type: str):
        """Add a bulleted or numbered list."""
        for item in items:
            if list_type == 'numbered':
                paragraph = self.document.add_paragraph(style='List Number')
            else:
                paragraph = self.document.add_paragraph(style='List Bullet')
            self._add_formatted_text(paragraph, item)
    
    def _add_blockquote(self, text: str):
        """Add a styled blockquote."""
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        
        # Add left border effect using background
        run = paragraph.add_run('│ ')
        run.font.color.rgb = RGBColor(99, 102, 241)  # Indigo
        run.bold = True
        
        self._add_formatted_text(paragraph, text)
    
    def _add_diagram_placeholder(self, description: str):
        """Add a placeholder for diagram description."""
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create a styled box for diagram
        run = paragraph.add_run('📊 DIAGRAM')
        run.bold = True
        run.font.color.rgb = RGBColor(16, 185, 129)  # Emerald
        
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.right_indent = Inches(0.5)
        
        run = paragraph.add_run(description)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
    
    def _add_table_from_markdown(self, table_lines: list):
        """Parse markdown table and add to document."""
        if not table_lines:
            return
        
        # Parse rows
        rows = []
        for line in table_lines:
            # Skip separator lines (----)
            if re.match(r'^[\s|:-]+$', line):
                continue
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
        
        if not rows:
            return
        
        # Determine number of columns
        num_cols = max(len(row) for row in rows)
        
        # Create table
        table = self.document.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        # Fill in cells
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    # First row is header
                    if i == 0:
                        cell.paragraphs[0].add_run(cell_text).bold = True
                    else:
                        cell.paragraphs[0].add_run(cell_text)
        
        # Add spacing after table
        self.document.add_paragraph()


def generate_docx(extracted_text: str, title: str = "Extracted Document") -> BytesIO:
    """
    Convenience function to generate a Word document.
    
    Args:
        extracted_text: Text with formatting markers
        title: Document title
        
    Returns:
        BytesIO object containing the .docx file
    """
    generator = DocxGenerator()
    return generator.create_document(extracted_text, title)
